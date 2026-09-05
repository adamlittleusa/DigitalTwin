"""Extract the paragraphs of a .docx into a markdown text file. One-off helper.

Usage: uv run --with python-docx python scripts/extract_docx.py <input.docx> <output.md>
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx


def extract(source: Path) -> str:
    document = docx.Document(str(source))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    rows = [
        " | ".join(cell.text.strip() for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    return "\n\n".join([*paragraphs, *rows]) + "\n"


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        sys.stderr.write(__doc__ or "")
        return 2
    source, target = Path(argv[1]), Path(argv[2])
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(extract(source), encoding="utf-8")
    sys.stderr.write(f"Wrote {target} ({target.stat().st_size} bytes)\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
